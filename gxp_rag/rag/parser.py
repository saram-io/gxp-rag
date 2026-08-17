"""GxP document parser supporting PDF, DOCX, Markdown, Text, and JSON formats."""

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from pydantic import BaseModel, Field

from gxp_rag.schemas.document import DocumentType


class ParsedGxPDocument(BaseModel):
    """Normalized parsed GxP document."""
    doc_id: str = Field(..., description="Document identifier (e.g., SOP-LAB-012)")
    title: str = Field(..., description="Title of the document")
    doc_type: DocumentType = Field(default=DocumentType.SOP)
    department: str = Field(default="Quality Assurance")
    version: str = Field(default="1.0")
    effective_date: Optional[str] = None
    raw_text: str
    metadata: Dict[str, Any] = Field(default_factory=dict)
    file_path: Optional[str] = None


class GxPDocumentParser:
    """Multi-format parser for GxP documents."""

    @staticmethod
    def detect_doc_type(text: str, filename: str) -> DocumentType:
        """Heuristically infer the GxP document type from content or filename."""
        text_lower = (text[:2000] + " " + filename).lower()
        
        if "deviation" in text_lower or "non-conformance" in text_lower or "ncr" in text_lower:
            return DocumentType.DEVIATION_REPORT
        elif "capa" in text_lower or "corrective action" in text_lower:
            return DocumentType.CAPA
        elif "work instruction" in text_lower or "wi-" in text_lower or "work-instruction" in text_lower:
            return DocumentType.WORK_INSTRUCTION
        elif "validation" in text_lower or "iq/oq" in text_lower or "qualification" in text_lower or "protocol" in text_lower:
            return DocumentType.VALIDATION_PROTOCOL
        elif "change control" in text_lower or "ccr" in text_lower:
            return DocumentType.CHANGE_CONTROL
        elif "batch record" in text_lower or "bpr" in text_lower or "master batch" in text_lower:
            return DocumentType.BATCH_PRODUCTION_RECORD
        elif "test method" in text_lower or "analytical" in text_lower or "specification" in text_lower:
            return DocumentType.ANALYTICAL_TEST_METHOD
        elif "guideline" in text_lower or "regulation" in text_lower:
            return DocumentType.REGULATORY_GUIDELINE
        else:
            return DocumentType.SOP

    @staticmethod
    def extract_doc_id(text: str, filename: str) -> str:
        """Extract Document ID from text headers or filename."""
        # Check standard GxP ID patterns: SOP-XYZ-123, DEV-2024-001, VAL-PRT-04, etc.
        id_patterns = [
            r"(?:Doc(?:ument)?\s*(?:ID|Number|No\.?)|SOP\s*No\.?)\s*[:#]?\s*([A-Z0-9_\-]+)",
            r"\b([A-Z]{2,5}-[A-Z0-9]{2,5}-\d{2,5})\b",
            r"\b(SOP-[A-Z0-9_\-]+)\b",
            r"\b(DEV-[A-Z0-9_\-]+)\b",
            r"\b(VAL-[A-Z0-9_\-]+)\b",
            r"\b(WI-[A-Z0-9_\-]+)\b",
            r"\b(CAPA-[A-Z0-9_\-]+)\b",
        ]
        for pattern in id_patterns:
            match = re.search(pattern, text[:1500], re.IGNORECASE)
            if match:
                return match.group(1).upper()
        
        # Fallback to sanitized stem of filename
        clean_stem = Path(filename).stem.upper().replace(" ", "-").replace("_", "-")
        return clean_stem or "DOC-GXP-001"

    @staticmethod
    def extract_title(text: str, filename: str) -> str:
        """Extract Title from first heading or document header."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines[:10]:
            if line.startswith("#"):
                clean = line.lstrip("#").strip()
                # If it looks like 'SOP-123: Title', strip the prefix
                if ":" in clean:
                    clean = clean.split(":", 1)[1].strip()
                if len(clean) > 3:
                    return clean
            elif line.lower().startswith("title:") or line.lower().startswith("subject:"):
                return line.split(":", 1)[1].strip()
        
        # Fallback to cleaned filename
        return Path(filename).stem.replace("_", " ").replace("-", " ").title()

    @staticmethod
    def extract_version(text: str) -> str:
        """Extract version string from document header."""
        match = re.search(r"(?:Version|Ver\.?|Rev\.?)\s*[:#]?\s*(\d+(?:\.\d+)?)", text[:1500], re.IGNORECASE)
        if match:
            return match.group(1)
        return "1.0"

    @staticmethod
    def extract_department(text: str) -> str:
        """Extract department from header if available."""
        match = re.search(r"(?:Department|Dept\.?)\s*[:#]?\s*([A-Za-z0-9 &/\-]+)", text[:1500], re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return "Quality Assurance & Operations"

    @classmethod
    def parse_file(cls, file_path: Union[str, Path]) -> ParsedGxPDocument:
        """Parse a document from file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix in [".md", ".markdown", ".txt"]:
            raw_text = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".pdf":
            raw_text = cls._parse_pdf(path)
        elif suffix in [".docx", ".doc"]:
            raw_text = cls._parse_docx(path)
        elif suffix == ".json":
            raw_text = cls._parse_json(path)
        else:
            raw_text = path.read_text(encoding="utf-8", errors="replace")

        doc_type = cls.detect_doc_type(raw_text, path.name)
        doc_id = cls.extract_doc_id(raw_text, path.name)
        title = cls.extract_title(raw_text, path.name)
        version = cls.extract_version(raw_text)
        department = cls.extract_department(raw_text)

        return ParsedGxPDocument(
            doc_id=doc_id,
            title=title,
            doc_type=doc_type,
            department=department,
            version=version,
            raw_text=raw_text,
            file_path=str(path.resolve()),
            metadata={
                "source_file": path.name,
                "file_size_bytes": path.stat().st_size,
            }
        )

    @classmethod
    def parse_text(
        cls,
        text: str,
        filename: str = "document.txt",
        doc_id: Optional[str] = None,
        title: Optional[str] = None,
        doc_type: Optional[DocumentType] = None,
        department: Optional[str] = None,
        version: Optional[str] = None,
    ) -> ParsedGxPDocument:
        """Parse in-memory raw text."""
        inferred_doc_type = doc_type or cls.detect_doc_type(text, filename)
        inferred_doc_id = doc_id or cls.extract_doc_id(text, filename)
        inferred_title = title or cls.extract_title(text, filename)
        inferred_version = version or cls.extract_version(text)
        inferred_department = department or cls.extract_department(text)

        return ParsedGxPDocument(
            doc_id=inferred_doc_id,
            title=inferred_title,
            doc_type=inferred_doc_type,
            department=inferred_department,
            version=inferred_version,
            raw_text=text,
            file_path=None,
            metadata={"source_file": filename}
        )

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Extract text from PDF using pypdf."""
        import pypdf
        reader = pypdf.PdfReader(str(path))
        extracted_pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            extracted_pages.append(f"--- Page {i+1} ---\n{page_text}")
        return "\n\n".join(extracted_pages)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def _parse_json(path: Path) -> str:
        """Parse structured JSON file."""
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            return json.dumps(data, indent=2)
        return str(data)
